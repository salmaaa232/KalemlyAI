import os
import io
import re
import sys
import shutil
import tempfile
from typing import List, Dict, Any

# Force UTF-8 stdout/stderr on Windows to avoid UnicodeEncodeError with special chars in PDF content
if hasattr(sys.stdout, 'reconfigure'):
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')
    except Exception:
        pass

try:
    import pypdf
except Exception:
    pypdf = None

try:
    import docx
except Exception:
    docx = None

try:
    from langchain_community.document_loaders import PyPDFLoader, TextLoader, WebBaseLoader
except Exception:
    PyPDFLoader = None
    TextLoader = None
    WebBaseLoader = None

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except Exception:
    RecursiveCharacterTextSplitter = None

try:
    from langchain_chroma import Chroma
except Exception:
    try:
        from langchain_community.vectorstores import Chroma
    except Exception:
        Chroma = None

try:
    from langchain_huggingface import HuggingFaceEmbeddings
except Exception:
    try:
        from langchain_community.embeddings import HuggingFaceEmbeddings
    except Exception:
        HuggingFaceEmbeddings = None

from langchain_core.documents import Document
from config import CHROMA_PERSIST_DIR

STOP_WORDS = {"the", "a", "an", "is", "are", "and", "or", "in", "on", "at", "to", "for", "of", "with", "you", "your", "my", "me", "what", "how", "why"}

SMALL_TALK_GREETINGS = {
    "hi", "hello", "hey", "good morning", "good evening", "good afternoon",
    "thanks", "thank you", "bye", "goodbye", "help", "who are you", "who r u",
    "مرحبا", "أهلا", "سلام", "شكرا", "مع السلامة"
}

def is_small_talk(query: str) -> bool:
    cleaned = re.sub(r'[^\w\s]', '', query.strip().lower())
    if cleaned in SMALL_TALK_GREETINGS:
        return True
    words = cleaned.split()
    if len(words) <= 2 and all(w in SMALL_TALK_GREETINGS for w in words):
        return True
    return False

def clean_extracted_text(text: str) -> str:
    """Removes raw PDF binary header markers, control chars, and noise."""
    if not text:
        return ""
    if text.startswith("%PDF") or "%PDF-1." in text[:50]:
        text = re.sub(r'%PDF-\d\.\d[^\n]*', '', text)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    return text.strip()

class RAGEngine:
    def __init__(self, persist_directory: str = CHROMA_PERSIST_DIR):
        self.persist_directory = os.path.abspath(persist_directory)
        self.uploads_directory = os.path.abspath("./uploads")
        os.makedirs(self.persist_directory, exist_ok=True)
        os.makedirs(self.uploads_directory, exist_ok=True)
        self.memory_store: Dict[str, List[Document]] = {}
        self.embeddings = None

        if HuggingFaceEmbeddings:
            try:
                self.embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
            except Exception as e:
                print(f"[WARNING] Could not initialize HuggingFaceEmbeddings: {e}")

    def process_and_store_documents(
        self,
        chatbot_id: str,
        files: List[Dict[str, Any]] = [],
        urls: List[str] = [],
        text_contents: List[str] = []
    ) -> int:
        documents: List[Document] = []
        
        # Save persistent copy of uploaded files to ./uploads/{chatbot_id}/
        bot_uploads_dir = os.path.join(self.uploads_directory, chatbot_id)
        os.makedirs(bot_uploads_dir, exist_ok=True)

        # 1. Files processing (PDF, DOCX, TXT)
        for file_obj in files:
            filename = file_obj.get("filename", "file.txt")
            content = file_obj.get("content", b"")

            if content and filename:
                try:
                    with open(os.path.join(bot_uploads_dir, filename), "wb") as f:
                        f.write(content)
                except Exception as save_err:
                    print(f"Warning saving raw file copy to disk: {save_err}")

            extracted_text = ""

            # PDF Parsing
            if filename.lower().endswith(".pdf") or content.startswith(b"%PDF"):
                if pypdf:
                    try:
                        reader = pypdf.PdfReader(io.BytesIO(content))
                        page_texts = []
                        for idx, page in enumerate(reader.pages):
                            t = page.extract_text()
                            if t and t.strip():
                                page_texts.append(f"[Page {idx+1}]\n{t.strip()}")
                        extracted_text = "\n".join(page_texts)
                    except Exception as e:
                        print(f"Error parsing PDF with pypdf {filename}: {e}")

            # DOCX Parsing
            elif filename.lower().endswith(".docx") or filename.lower().endswith(".doc"):
                if docx:
                    try:
                        doc = docx.Document(io.BytesIO(content))
                        extracted_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                    except Exception as e:
                        print(f"Error parsing DOCX {filename}: {e}")

            # TXT / Plain text fallback
            if not extracted_text.strip() and not filename.lower().endswith(".pdf"):
                try:
                    extracted_text = content.decode("utf-8", errors="ignore")
                except Exception as e:
                    print(f"Error decoding text file {filename}: {e}")

            cleaned = clean_extracted_text(extracted_text)
            if cleaned and not cleaned.startswith("%PDF"):
                documents.append(Document(page_content=cleaned, metadata={"source": filename}))

        # 2. URLs processing
        for url in urls:
            if not url.strip():
                continue
            if WebBaseLoader:
                try:
                    loader = WebBaseLoader(url.strip())
                    docs = loader.load()
                    for d in docs:
                        cleaned = clean_extracted_text(d.page_content)
                        if cleaned:
                            documents.append(Document(page_content=cleaned, metadata={"source": url}))
                    continue
                except Exception:
                    pass
            documents.append(Document(page_content=f"Website reference: {url}", metadata={"source": url}))

        # 3. Text contents (Business description & Custom Instructions)
        for idx, text in enumerate(text_contents):
            cleaned = clean_extracted_text(text)
            if cleaned:
                documents.append(Document(page_content=cleaned, metadata={"source": f"custom_instructions_{idx+1}"}))

        if not documents:
            return 0

        # Chunking
        if RecursiveCharacterTextSplitter:
            try:
                splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=100)
                chunks = splitter.split_documents(documents)
            except Exception:
                chunks = documents
        else:
            chunks = documents

        # Store in RAM memory store
        self.memory_store[chatbot_id] = chunks

        # Persist to disk via ChromaDB
        if Chroma and self.embeddings:
            try:
                chatbot_dir = os.path.join(self.persist_directory, chatbot_id)
                if os.path.exists(chatbot_dir):
                    shutil.rmtree(chatbot_dir, ignore_errors=True)
                Chroma.from_documents(documents=chunks, embedding=self.embeddings, persist_directory=chatbot_dir)
                print(f"[OK] Saved ChromaDB vector store to disk for chatbot '{chatbot_id}' at: {chatbot_dir}")
            except Exception as e:
                print(f"[WARNING] Chroma storage warning: {e}")

        return len(chunks)

    def retrieve_chunk_objects(self, chatbot_id: str, query: str, k: int = 3) -> List[Dict[str, str]]:
        chatbot_dir = os.path.join(self.persist_directory, chatbot_id)
        disk_exists = os.path.exists(chatbot_dir)
        is_small = is_small_talk(query)
        results: List[Dict[str, str]] = []
        seen_contents = set()
        status_reason = ""

        # 1. Do not retrieve chunks for small talk or greetings
        if is_small:
            status_reason = "Small talk greeting detected; vector retrieval skipped."
            self._log_debug(chatbot_id, query, disk_exists, 0, status_reason, [])
            return []

        # 2. Auto-recovery helper: Reload vector store from disk uploads & DB if missing
        if not disk_exists:
            try:
                bot_uploads_dir = os.path.join(self.uploads_directory, chatbot_id)
                recovered_files = []
                if os.path.exists(bot_uploads_dir):
                    for fname in os.listdir(bot_uploads_dir):
                        fpath = os.path.join(bot_uploads_dir, fname)
                        if os.path.isfile(fpath):
                            with open(fpath, "rb") as f:
                                recovered_files.append({"filename": fname, "content": f.read()})
                
                from database import get_chatbot
                bot = get_chatbot(chatbot_id)
                text_contents = []
                if bot:
                    if bot.get("business_description"):
                        text_contents.append(f"Business Overview for {bot.get('company_name')}:\n{bot.get('business_description')}")
                    if bot.get("instructions"):
                        text_contents.append(f"Chatbot Rules & Custom Instructions:\n{bot.get('instructions')}")
                        
                if recovered_files or text_contents:
                    print(f"[AUTO RECOVERY] Re-indexing missing Chroma vector store for bot '{chatbot_id}' ({len(recovered_files)} files)...")
                    self.process_and_store_documents(
                        chatbot_id=chatbot_id,
                        files=recovered_files,
                        text_contents=text_contents
                    )
                    disk_exists = os.path.exists(chatbot_dir)
            except Exception as recovery_err:
                print(f"[RECOVERY WARNING] Could not auto-recover vector store: {recovery_err}")

        # 3. Disk Retrieval via ChromaDB
        if Chroma and self.embeddings:
            if disk_exists:
                try:
                    vectorstore = Chroma(persist_directory=chatbot_dir, embedding_function=self.embeddings)
                    docs = vectorstore.similarity_search(query, k=k)
                    for d in docs:
                        cleaned = d.page_content.strip()
                        if cleaned and not cleaned.startswith("%PDF") and cleaned not in seen_contents:
                            seen_contents.add(cleaned)
                            results.append({
                                "source": d.metadata.get("source", "Document"),
                                "content": cleaned
                            })
                    if results:
                        status_reason = "Retrieval successful from ChromaDB disk store."
                        self._log_debug(chatbot_id, query, disk_exists, len(results[:k]), status_reason, results[:k])
                        return results[:k]
                    else:
                        status_reason = "ChromaDB vector store searched on disk, but 0 matching chunks found."
                except Exception as e:
                    status_reason = f"ChromaDB disk loading exception: {e}"
            else:
                status_reason = f"Vector store directory missing on disk at '{chatbot_dir}'."
        else:
            status_reason = "Embeddings model or Chroma library not initialized."

        # 4. Fallback to RAM Memory Store
        docs = self.memory_store.get(chatbot_id, [])
        if docs:
            query_terms = [t for t in query.lower().split() if len(t) > 2 and t not in STOP_WORDS]
            scored_docs = []
            for d in docs:
                content_lower = d.page_content.lower()
                if not content_lower.startswith("%PDF"):
                    score = sum(1 for term in query_terms if term in content_lower)
                    if score > 0:
                        scored_docs.append((score, d))

            scored_docs.sort(key=lambda x: x[0], reverse=True)
            for score, d in scored_docs:
                cleaned = d.page_content.strip()
                if cleaned not in seen_contents:
                    seen_contents.add(cleaned)
                    results.append({
                        "source": d.metadata.get("source", "Document"),
                        "content": cleaned
                    })
                if len(results) >= k:
                    break

            if results:
                status_reason += " (Retrieved chunks from RAM memory store fallback)."

        self._log_debug(chatbot_id, query, disk_exists, len(results[:k]), status_reason, results[:k])
        return results[:k]


    def _safe_print(self, text: str):
        """Print safely, replacing any characters that can't be encoded."""
        try:
            print(text)
        except UnicodeEncodeError:
            print(text.encode('ascii', errors='replace').decode('ascii'))

    def _log_debug(self, chatbot_id: str, query: str, disk_exists: bool, count: int, reason: str, chunks: List[Dict[str, str]]):
        try:
            self._safe_print("\n" + "="*60)
            self._safe_print("[RAG DIAGNOSTIC]")
            self._safe_print(f"  Chatbot ID:            {chatbot_id}")
            self._safe_print(f"  Query:                 '{query}'")
            self._safe_print(f"  Chroma Directory:      {os.path.join(self.persist_directory, chatbot_id)}")
            self._safe_print(f"  Disk Store Exists:     {disk_exists}")
            self._safe_print(f"  Embeddings Ready:      {self.embeddings is not None}")
            self._safe_print(f"  Retrieved Chunks:      {count}")
            self._safe_print(f"  Reason / Status:       {reason}")
            if chunks:
                self._safe_print("  --------------------------------------------------")
                for idx, chk in enumerate(chunks, 1):
                    snip = chk['content'][:80].replace('\n', ' ')
                    self._safe_print(f"  Chunk #{idx} [{chk['source']}]: {snip}...")
            self._safe_print("="*60 + "\n")
        except Exception:
            pass  # Never let logging crash the retrieval pipeline

    def retrieve_context(self, chatbot_id: str, query: str, k: int = 3) -> str:
        chunks = self.retrieve_chunk_objects(chatbot_id, query, k=k)
        if not chunks:
            return ""
        return "\n\n".join([f"[Source: {c['source']}]\n{c['content']}" for c in chunks])

rag_engine = RAGEngine()
